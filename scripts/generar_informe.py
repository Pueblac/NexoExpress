import os
import glob
import datetime
import base64
import requests
from docx import Document
from docx.shared import Inches

def renderizar_mermaid_a_png(mmd_path, png_path):
    print(f"Renderizando {mmd_path}...")
    with open(mmd_path, 'r', encoding='utf-8') as f:
        graph = f.read()
    
    # Codificamos el texto mermaid a base64
    graphbytes = graph.encode("utf-8")
    base64_bytes = base64.b64encode(graphbytes)
    base64_string = base64_bytes.decode("utf-8")
    
    # Llamamos a la API de mermaid.ink
    url = f"https://mermaid.ink/img/{base64_string}"
    response = requests.get(url)
    
    if response.status_code == 200:
        with open(png_path, 'wb') as f:
            f.write(response.content)
        print(f"✅ Imagen generada: {png_path}")
    else:
        print(f"❌ Error al generar imagen para {mmd_path}. Status: {response.status_code}")

def generar_informe_docx(md_path, output_docx_path, flujos_dir):
    print("Creando plantilla del documento Word...")
    doc = Document()
    doc.add_heading('Informe de Arquitectura - NexoExpress', 0)

    # 1. Leer el markdown del Arquitecto y añadirlo al documento
    if os.path.exists(md_path):
        print("Añadiendo texto del análisis del Arquitecto...")
        with open(md_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line:
                    doc.add_paragraph(line)
    else:
        print(f"Advertencia: No se encontró el informe {md_path}")
        doc.add_paragraph("No se encontró el reporte del arquitecto.")

    # 2. Añadir los diagramas renderizados (PNGs) al final
    doc.add_heading('Anexos: Diagramas de Flujo y Arquitectura', level=1)
    png_files = glob.glob(os.path.join(flujos_dir, "*.png"))
    
    if not png_files:
        doc.add_paragraph("No se encontraron diagramas de flujo renderizados.")
    
    for png_file in png_files:
        nombre = os.path.basename(png_file).replace(".png", "")
        doc.add_heading(nombre.replace("_", " ").title(), level=2)
        try:
            # Añadir la imagen con un ancho máximo para que quepa en la hoja
            doc.add_picture(png_file, width=Inches(6.0))
        except Exception as e:
            doc.add_paragraph(f"[Error al insertar imagen {nombre}: {e}]")

    doc.save(output_docx_path)
    print(f"🚀 ¡Informe generado exitosamente en {output_docx_path}!")

if __name__ == "__main__":
    # Definición de rutas
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    flujos_dir = os.path.join(base_dir, "docs", "flujos")
    informes_dir = os.path.join(base_dir, "docs", "informes")
    md_file = os.path.join(base_dir, "docs", "informe_arquitectura.md")
    
    timestamp = datetime.datetime.now().strftime("%d_%m_%y_%H_%M")
    out_docx = os.path.join(informes_dir, f"informe_consolidado_{timestamp}.docx")

    # Asegurar que existan los directorios
    os.makedirs(flujos_dir, exist_ok=True)
    os.makedirs(informes_dir, exist_ok=True)

    # Convertir todos los .mmd nuevos a .png
    mmd_files = glob.glob(os.path.join(flujos_dir, "*.mmd"))
    for mmd_file in mmd_files:
        png_file = mmd_file.replace('.mmd', '.png')
        # Renderizar si el PNG no existe o el MMD fue modificado más recientemente
        if not os.path.exists(png_file) or os.path.getmtime(mmd_file) > os.path.getmtime(png_file):
            renderizar_mermaid_a_png(mmd_file, png_file)

    # Finalmente, ensamblar el documento Word
    generar_informe_docx(md_file, out_docx, flujos_dir)
